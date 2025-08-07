import requests
import json
import webview
import io
import base64
import socket
from googletrans import Translator
from flask import Flask, request, jsonify, send_from_directory

from transformers import pipeline, AutoTokenizer, GPT2LMHeadModel, GPT2Tokenizer
from textblob import TextBlob
import logging
import torch
import rispy
import bibtexparser
import mammoth
import base64
import io
import re
import concurrent.futures
from docx import Document
from flask import Flask, request, jsonify, send_from_directory
app = Flask(__name__)

def find_available_port(start_port=5000):
    """Find an available port starting from start_port"""
    for port in range(start_port, start_port + 100):
        try:
            sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            result = sock.connect_ex(('127.0.0.1', port))
            sock.close()
            if result != 0:  # Port is available
                return port
        except:
            continue
    return start_port  # Fallback
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
try:
    import textract
except ImportError:
    textract = None

# AI pipeline gratis dari HuggingFace (misal T5) - LAZY LOADING
nlp = None
lang_detector = None
tokenizer = None

def init_models():
    global nlp, lang_detector, tokenizer
    if nlp is None:
        print("Loading NLP models...")
        nlp = pipeline("text2text-generation", model="google/flan-t5-base")
        lang_detector = pipeline("text-classification", model="papluca/xlm-roberta-base-language-detection")
        tokenizer = AutoTokenizer.from_pretrained("google/flan-t5-base")
        print("Models loaded successfully!")

def mistral_fix_typo(chunk):
    url = "https://openrouter.ai/api/v1/chat/completions"
    api_key = "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"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    prompt = (
        "Perbaiki hanya typo dan ejaan pada teks berikut. "
        "JANGAN menambah, mengubah, meringkas, atau menghilangkan kalimat dan isi. "
        "JANGAN menambahkan penjelasan, interpretasi, daftar perubahan, atau format lain. "
        "Kembalikan hasil persis seperti input, hanya dengan typo/ejaan yang sudah diperbaiki. "
        "JANGAN tambahkan apapun selain teks hasil perbaikan.\n\n"
        f"{chunk}"
    )
    data = {
        "model": "mistralai/mistral-small-3.2-24b-instruct:free",
        "messages": [
            {
                "role": "user",
                "content": [
                    { "type": "text", "text": prompt }
                ]
            }
        ]
    }
    try:
        response = requests.post(url, headers=headers, data=json.dumps(data), timeout=60)
        result = response.json()
        return result.get("choices", [{}])[0].get("message", {}).get("content", "")
    except Exception as e:
        return f"[Gagal AI Mistral: {str(e)}]"

def chunk_by_tokens(text, max_tokens=384):
    # Inisialisasi model jika belum di-load
    if tokenizer is None:
        init_models()
        
    words = text.split()
    chunks = []
    current_chunk = []
    current_len = 0
    
    for word in words:
        tokens = tokenizer.tokenize(word)
        if current_len + len(tokens) > max_tokens:
            chunks.append(' '.join(current_chunk))
            current_chunk = [word]
            current_len = len(tokens)
        else:
            current_chunk.append(word)
            current_len += len(tokens)
    
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    
    return chunks

def detect_language(text):
    # Inisialisasi model jika belum di-load
    if lang_detector is None or tokenizer is None:
        init_models()
    
    # Potong teks jika terlalu panjang
    words = text.split()
    selected_words = []
    current_len = 0
    
    for word in words:
        tokens = tokenizer.tokenize(word)
        if current_len + len(tokens) > 384:  # Batas yang lebih rendah
            break
        selected_words.append(word)
        current_len += len(tokens)
    
    text = ' '.join(selected_words)
    result = lang_detector(text)
    return result[0]['label']

# Helper untuk angka romawi
def to_roman(num):
    val = [
        1000, 900, 500, 400,
        100, 90, 50, 40,
        10, 9, 5, 4, 1
    ]
    syb = [
        "M", "CM", "D", "CD",
        "C", "XC", "L", "XL",
        "X", "IX", "V", "IV", "I"
    ]
    roman_num = ''
    i = 0
    while num > 0:
        for _ in range(num // val[i]):
            roman_num += syb[i]
            num -= val[i]
        i += 1
    return roman_num
    
class Api:
    def format_cleanup(self, text, mode):
        def process_chunk(chunk, mode):
            lang = detect_language(chunk)
            if mode == "fix_typo":
                if lang == "en":
                    import language_tool_python
                    tool = language_tool_python.LanguageTool('en-US', host='localhost')
                    matches = tool.check(chunk)
                    return language_tool_python.utils.correct(chunk, matches)
                else:
                    return mistral_fix_typo(chunk)
            elif mode == "capitalize_sentence":
                prompt = f"Ubah huruf pertama setiap kalimat setelah titik menjadi kapital:\n{chunk}"
                return mistral_fix_typo(prompt)
            elif mode == "tab_sentence":
                prompt = f"Tambahkan tab di awal setiap kalimat:\n{chunk}"
                return mistral_fix_typo(prompt)
            elif mode == "auto_format":
                prompt = (
                    "Rapikan seluruh format penulisan dokumen berikut (ejaan, typo, tab, spasi) tanpa mengubah struktur, urutan, atau isi. "
                    "Jangan menambah penjelasan atau interpretasi tambahan.\n"
                    f"{chunk}"
                )
                return mistral_fix_typo(prompt)
            else:
                return chunk

        if mode in ["capitalize_sentence", "tab_sentence", "auto_format", "fix_typo"]:
            chunks = chunk_by_tokens(text, max_tokens=384)
            results = []
            for chunk in chunks:
                if chunk.strip():
                    processed = process_chunk(chunk, mode)
                    results.append(processed)
            result = ' '.join(results)
        elif mode == "fix_spacing":
            import re
            result = re.sub(r'\s+', ' ', text).strip()
        else:
            result = text

        return result

    def translate(self, text, src, dest):
        # Validasi kode bahasa
        valid_langs = ['id', 'en', 'ja', 'zh-CN', 'ar', 'es', 'fr', 'de', 'ru', 'ko', 'it', 'pt', 'auto']
        print(f"[DEBUG] translate src={src}, dest={dest}")
        if src not in valid_langs:
            src = 'auto'
        if dest == 'auto':
            return "Bahasa tujuan tidak boleh 'auto'. Harap pilih secara eksplisit."
        elif dest not in valid_langs:
            return f"Bahasa tujuan '{dest}' tidak valid. Pilihan: {', '.join([l for l in valid_langs if l != 'auto'])}"
        try:
            translator = Translator()
            max_len = 4000
            parts = []
            current = ''
            for paragraph in text.split('\n'):
                while len(paragraph) > max_len:
                    split_pos = paragraph.rfind(' ', 0, max_len)
                    if split_pos == -1:
                        split_pos = max_len
                    parts.append(paragraph[:split_pos])
                    paragraph = paragraph[split_pos:]
                current += paragraph + '\n'
                if len(current) >= max_len:
                    split_pos = current.rfind(' ', 0, max_len)
                    if split_pos == -1:
                        split_pos = max_len
                    parts.append(current[:split_pos])
                    current = current[split_pos:]
            if current.strip():
                parts.append(current)
            # Translate tiap bagian
            translated = []
            for part in parts:
                result = translator.translate(part, src=src, dest=dest)
                translated.append(result.text)
            return '\n'.join(translated)
        except Exception as e:
            return f"Error: {str(e)}"   

    def mistral_translate(self, text, src="id", dest="en"):
        import requests
        import json

        max_len = 4000
        parts = []
        current = ''
        for paragraph in text.split('\n'):
            while len(paragraph) > max_len:
                split_pos = paragraph.rfind(' ', 0, max_len)
                if split_pos == -1:
                    split_pos = max_len
                parts.append(paragraph[:split_pos])
                paragraph = paragraph[split_pos:]
            current += paragraph + '\n'
            if len(current) >= max_len:
                split_pos = current.rfind(' ', 0, max_len)
                if split_pos == -1:
                    split_pos = max_len
                parts.append(current[:split_pos])
                current = current[split_pos:]
        if current.strip():
            parts.append(current)

        def single_mistral_translate(part, src, dest):
            url = "https://openrouter.ai/api/v1/chat/completions"
            api_key = "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"
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            }
            prompt = f"Terjemahkan teks berikut dari {src} ke {dest} secara akurat dan natural:\n{part}"
            data = {
                "model": "mistralai/mistral-small-3.2-24b-instruct:free",
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            { "type": "text", "text": prompt }
                        ]
                    }
                ]
            }
            try:
                response = requests.post(url, headers=headers, data=json.dumps(data))
                result = response.json()
                return result.get("choices", [{}])[0].get("message", {}).get("content", "")
            except Exception as e:
                return f"[Gagal translate Mistral: {str(e)}]"

        translated = []
        for part in parts:
            translated.append(single_mistral_translate(part, src, dest))
        return '\n'.join(translated)
    
    # Removed duplicate chat_llama method
    def chat_llama(self, message, history_json):
        import requests, json, traceback
        print("============= CHAT LLAMA CALLED =============")
        print(f"Received message: {message}")
        print(f"Received history: {history_json}")
        url = "https://openrouter.ai/api/v1/chat/completions"
        api_key = "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"  # API key model 1
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        }
        try:
            history = json.loads(history_json)
            print(f"Parsed history: {history}")
            system_message = {
                "role": "system",
                "content": (
                    "Jawablah sebagai asisten dokumen profesional. "
                    "Balas dengan gaya ramah, gunakan emoji jika relevan dengan nuansa topik, dan tampilkan jawaban dalam format chat yang rapi. "
                    "Jangan gunakan tanda bintang (**), markdown, atau format lain yang tidak natural untuk chat. "
                    "Pisahkan poin-poin dengan baris baru, dan gunakan bahasa Indonesia yang mudah dipahami."
                )
            }
            messages = [system_message] + history
            data = {
                "model": "meta-llama/llama-3.3-70b-instruct:free",
                "messages": messages
            }
            print(f"Sending request to OpenRouter with data: {data}")
            response = requests.post(url, headers=headers, data=json.dumps(data), timeout=60)
            print("Response status code:", response.status_code)
            print("Response content:", response.text)
            result = response.json()
            print("OpenRouter raw response:", result)  # DEBUG LOG
            if "error" in result:
                error_msg = f"Maaf, terjadi error pada AI: {result['error'].get('message', 'Unknown error')}"
                print(f"Error in response: {error_msg}")
                return error_msg
            reply = result.get("choices", [{}])[0].get("message", {}).get("content", "")
            print(f"AI reply: {reply}")
            if not reply or not reply.strip():
                reply = "Maaf, AI tidak dapat memberikan jawaban saat ini. Silakan coba lagi nanti."
            return reply
        except Exception as e:
            error_msg = f"Maaf, terjadi error pada AI: {str(e)}"
            print(f"Exception: {error_msg}")
            print(f"Exception details: {type(e).__name__} - {str(e)}")
            print(traceback.format_exc())
            return error_msg

    def read_file(self, file_b64, ext):
        data = base64.b64decode(file_b64)
        ext = ext.lower()
        if ext == 'docx':
            file_stream = io.BytesIO(data)
            result = mammoth.extract_raw_text(file_stream)
            html = result.value.replace('<!-- PAGE BREAK -->', '<hr class="page-separator"><div class="page-separator-label">Page</div>')
            return html
        elif ext == 'pdf' and PyPDF2:
            try:
                file_stream = io.BytesIO(data)
                reader = PyPDF2.PdfReader(file_stream)
                text = ''
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + '\n'
                if not text.strip():
                    return 'Tidak ada teks yang dapat diekstrak dari PDF.'
                return text
            except Exception as e:
                return f'Gagal membaca PDF: {str(e)}'
        elif ext == 'doc' and textract:
            # textract mendukung doc, docx, pdf, dll
            file_stream = io.BytesIO(data)
            text = textract.process('', input_stream=file_stream, extension='doc').decode('utf-8')
            return text
        else:
            return 'Format file tidak didukung atau library belum terpasang.'

    def auto_format(self, text):
        url = "https://openrouter.ai/api/v1/chat/completions"
        api_key = "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"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        }
        prompt = (
            "Rapikan struktur penulisan dokumen berikut sesuai standar akademik. "
            "Pastikan setiap paragraf memiliki indentasi, heading bab/subbab konsisten, dan spasi antar bagian proporsional. "
            "Perbaiki tata bahasa, ejaan, dan format penulisan agar lebih rapi, tetapi jangan mengubah isi, urutan, atau struktur dokumen. "
            "Jika ada daftar pustaka, kutipan, tabel, atau gambar, formatkan sesuai standar jurnal. "
            "Tambahkan satu baris kosong sebelum dan sesudah berupa spasi bagian penjelasan hasil auto format agar hasil dan penjelasan mudah dibaca.\n"
            "Kembalikan hasil dalam format HTML yang siap untuk editor dokumen.\n"
            "Jangan melakukan interpretasi tambahan di luar instruksi.\n\n"
            f"Dokumen:\n{text}"
        )
        data = {
            "model": "mistralai/mistral-small-3.2-24b-instruct:free",
            "messages": [
                {
                    "role": "user",
                    "content": [
                        { "type": "text", "text": prompt }
                    ]
                }
            ]
        }
        response = requests.post(url, headers=headers, data=json.dumps(data))
        result = response.json()
        return result.get("choices", [{}])[0].get("message", {}).get("content", "Gagal auto format")

    def apply_numbering(self, file_b64, numbering_type, from_page, to_page, position):
        data = base64.b64decode(file_b64)
        file_stream = io.BytesIO(data)
        doc = Document(file_stream)
        preview = ""
        page_num = 1

        for i, p in enumerate(doc.paragraphs, start=1):
            # Proses penomoran
            if from_page <= i <= to_page:
                if numbering_type == "regular":
                    p.text = f"{i} - {p.text}"
                elif numbering_type == "lower-roman":
                    p.text = f"{to_roman(i).lower()} - {p.text}"
                elif numbering_type == "upper-roman":
                    p.text = f"{to_roman(i)} - {p.text}"
                elif numbering_type == "combined":
                    p.text = f"{to_roman(i)}-{i} - {p.text}"

            # Tambahkan paragraph ke preview
            preview += f"{p.text}<br>"

            # Cek apakah ada page break di paragraph ini
            for run in p.runs:
                if 'page' in run._element.xml and '<w:br w:type="page"/>' in run._element.xml:
                    # Tambahkan pemisah halaman
                    preview += f'<hr class="page-separator"><div class="page-separator-label">Page {page_num+1}</div>'
                    page_num += 1

        return preview

    def download_numbered_doc(self, file_b64, numbering_type, from_page, to_page, position):
        import base64, io
        from docx import Document
        # Proses file dan penomoran
        data = base64.b64decode(file_b64)
        file_stream = io.BytesIO(data)
        doc = Document(file_stream)
        for i, p in enumerate(doc.paragraphs, start=1):
            if from_page <= i <= to_page:
                if numbering_type == "regular":
                    p.text = f"{i} - {p.text}"
                elif numbering_type == "lower-roman":
                    p.text = f"{to_roman(i).lower()} - {p.text}"
                elif numbering_type == "upper-roman":
                    p.text = f"{to_roman(i)} - {p.text}"
                elif numbering_type == "combined":
                    p.text = f"{to_roman(i)}-{i} - {p.text}"
        output_stream = io.BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)
        return base64.b64encode(output_stream.read()).decode('utf-8')
    
    def parse_reference(self, content, ext):
        ext = ext.lower()
        refs = []
        if ext in ['ris', 'txt']:
            # Parsing RIS
            for entry in content.split('\n\n'):
                if entry.strip():
                    refs.append({'raw': entry.strip()})
        elif ext == 'bib':
            # Parsing BibTeX
            for entry in content.split('@'):
                if entry.strip():
                    refs.append({'raw': '@' + entry.strip()})
        elif ext in ['doc', 'docx']:
            # Parsing Word (simple, bisa pakai mammoth)
            import base64, io
            from docx import Document
            if content.startswith('data:'):
                b64 = content.split(',')[1]
                data = base64.b64decode(b64)
            else:
                data = content.encode()
            file_stream = io.BytesIO(data)
            doc = Document(file_stream)
            for p in doc.paragraphs:
                if p.text.strip():
                    refs.append({'raw': p.text.strip()})
        # TODO: parsing EndNote, RefMan, RefWorks jika format diketahui
        return refs

    def format_reference(self, ref, style="APA"):
        def part(label, value, prefix='', suffix=''):
            return f"{prefix}{value}{suffix}" if value and value != 'N/A' else ''
        if style == "APA":
            return (
                f"{part('', ref.get('author'))}"
                f"{part(' (', ref.get('year'), '(', ')')}. "
                f"{part('', ref.get('title'))}. "
                f"{part('', ref.get('journal'))}"
                f"{part(', ', ref.get('volume'))}"
                f"{part('(', ref.get('issue'), '(', ')')}"
                f"{part(', ', ref.get('pages'))}"
            ).strip().replace('  ', ' ')
        elif style == "IEEE":
            return (
                f"{part('', ref.get('author'))}, "
                f"\"{part('', ref.get('title'))},\" "
                f"{part('', ref.get('journal'))}"
                f"{part(', vol. ', ref.get('volume'))}"
                f"{part(', no. ', ref.get('issue'))}"
                f"{part(', pp. ', ref.get('pages'))}"
                f"{part(', ', ref.get('year'))}."
            ).strip().replace('  ', ' ')
        elif style == "CHICAGO":
            return (
                f"{part('', ref.get('author'))}. "
                f"\"{part('', ref.get('title'))}.\" "
                f"{part('', ref.get('journal'))} "
                f"{part('', ref.get('volume'))}"
                f"{part(', no. ', ref.get('issue'))}"
                f"{part(' (', ref.get('year'), '(', ')')}: "
                f"{part('', ref.get('pages'))}."
            ).strip().replace('  ', ' ')
        elif style == "HARVARD":
            return (
                f"{part('', ref.get('author'))}, "
                f"{part('', ref.get('year'))}. "
                f"{part('', ref.get('title'))}. "
                f"{part('', ref.get('journal'))}"
                f"{part(', ', ref.get('volume'))}"
                f"{part('(', ref.get('issue'), '(', ')')}"
                f"{part(', pp. ', ref.get('pages'))}."
            ).strip().replace('  ', ' ')
        elif style == "MLA":
            return (
                f"{part('', ref.get('author'))}. "
                f"\"{part('', ref.get('title'))}.\" "
                f"{part('', ref.get('journal'))}"
                f"{part(', vol. ', ref.get('volume'))}"
                f"{part(', no. ', ref.get('issue'))}"
                f"{part(', ', ref.get('year'))}"
                f"{part(', pp. ', ref.get('pages'))}."
            ).strip().replace('  ', ' ')
        return ref.get('raw', '')

    def ai_parse_docx_reference(self, doc_text):
        url = "https://openrouter.ai/api/v1/chat/completions"
        api_key = "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"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        }
        prompt = (
            "Ekstrak daftar referensi berikut menjadi JSON array dengan field author, year, title, journal, volume, issue, pages. "
            "Gunakan format penulisan sitasi sesuai contoh di bawah ini untuk setiap style, tetapi JANGAN hasilkan contoh di bawah sebagai output. "
            "Jika ada informasi yang tidak diketahui (misal volume, issue, pages, dsb), jangan tulis field tersebut di hasil sitasi. "
            "Jangan ubah urutan, jangan tambahkan penjelasan. Hanya hasil JSON saja. "
            "Jangan singkat atau potong nama jurnal, judul, atau penulis. Tampilkan seluruh informasi yang tersedia.\n\n"
            "Contoh format sitasi (hanya untuk referensi, bukan untuk output):\n"
            "✅ 1. APA (American Psychological Association)\n"
            "Simatupang, H. B. (2019). Peranan perbankan dalam meningkatkan perekonomian Indonesia. JRAM (Jurnal Riset Akuntansi dan Manajemen), 9(2), 112–125.\n"
            "✅ 2. Chicago Style (Notes and Bibliography)\n"
            "Simatupang, H. B. \"Peranan Perbankan dalam Meningkatkan Perekonomian Indonesia.\" JRAM (Jurnal Riset Akuntansi dan Manajemen) 9, no. 2 (2019): 112–125.\n"
            "✅ 3. IEEE (Institute of Electrical and Electronics Engineers)\n"
            "[1] H. B. Simatupang, \"Peranan perbankan dalam meningkatkan perekonomian Indonesia,\" JRAM (Jurnal Riset Akuntansi dan Manajemen), vol. 9, no. 2, pp. 112–125, 2019.\n"
            "✅ 4. Harvard Style\n"
            "Simatupang, H.B., 2019. Peranan perbankan dalam meningkatkan perekonomian Indonesia. JRAM (Jurnal Riset Akuntansi dan Manajemen), 9(2), pp.112–125.\n"
            "✅ 5. MLA (Modern Language Association)\n"
            "Simatupang, H. B. \"Peranan Perbankan dalam Meningkatkan Perekonomian Indonesia.\" JRAM (Jurnal Riset Akuntansi dan Manajemen), vol. 9, no. 2, 2019, pp. 112–125.\n\n"
            f"{doc_text}"
        )
        data = {
            "model": "mistralai/mistral-small-3.2-24b-instruct:free",
            "messages": [
                {
                    "role": "user",
                    "content": [
                        { "type": "text", "text": prompt }
                    ]
                }
            ]
        }
        response = requests.post(url, headers=headers, data=json.dumps(data))
        result = response.json()
        # Ambil hasil JSON dari response
        import re, json as js
        match = re.search(r'\[.*\]', result.get("choices", [{}])[0].get("message", {}).get("content", ""))
        if match:
            try:
                refs = js.loads(match.group(0))
                return refs
            except Exception:
                pass
        return []

    def parse_reference(self, content, ext, style="APA"):
        ext = ext.lower()
        refs = []
        if ext == 'ris':
            entries = rispy.loads(content)
            for entry in entries:
                refs.append({
                    'author': ', '.join(entry.get('authors', [])),
                    'year': entry.get('year', ''),
                    'title': entry.get('title', ''),
                    'journal': entry.get('journal_name', ''),
                    'volume': entry.get('volume', ''),
                    'issue': entry.get('number', ''),
                    'pages': entry.get('start_page', '') + '-' + entry.get('end_page', ''),
                    'raw': entry.get('title', '')
                })
        elif ext == 'bib':
            bib_db = bibtexparser.loads(content)
            for entry in bib_db.entries:
                refs.append({
                    'author': entry.get('author', ''),
                    'year': entry.get('year', ''),
                    'title': entry.get('title', ''),
                    'journal': entry.get('journal', ''),
                    'volume': entry.get('volume', ''),
                    'issue': entry.get('number', ''),
                    'pages': entry.get('pages', ''),
                    'raw': entry.get('title', '')
                })
        elif ext in ['doc', 'docx']:
            import base64, io
            from docx import Document
            if content.startswith('data:'):
                b64 = content.split(',')[1]
                data = base64.b64decode(b64)
            else:
                data = content.encode()
            file_stream = io.BytesIO(data)
            doc = Document(file_stream)
            doc_text = '\n'.join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
            # Gunakan AI untuk parsing referensi
            refs = self.ai_parse_docx_reference(doc_text)
        
        # === VALIDASI FIELD KOSONG ===
        for ref in refs:
            for key in ['author', 'year', 'title', 'journal', 'volume', 'issue', 'pages']:
                if not ref.get(key):
                    ref[key] = 'N/A'
        # ...parsing format lain jika ada...
        formatted_refs = [self.format_reference(ref, style) for ref in refs]
        return [{'raw': r} for r in formatted_refs]

    def calculate_gltr_score(self, text, model_name="gpt2", chunk_size=512):
        tokenizer = GPT2Tokenizer.from_pretrained(model_name)
        model = GPT2LMHeadModel.from_pretrained(model_name)
        sentences = text.split('. ')
        chunks, current_chunk = [], ""
        for sentence in sentences:
            if len(current_chunk) + len(sentence) < chunk_size:
                current_chunk += sentence + ". "
            else:
                chunks.append(current_chunk)
                current_chunk = sentence + ". "
        if current_chunk:
            chunks.append(current_chunk)
        green_yellow, total = 0, 0
        for chunk in chunks:
            inputs = tokenizer(chunk, return_tensors="pt", truncation=True, max_length=1024)
            tokens = tokenizer.convert_ids_to_tokens(inputs["input_ids"][0])
            with torch.no_grad():
                logits = model(**inputs, labels=inputs["input_ids"]).logits[0]
            top_k = 1000
            predictions = torch.topk(logits, top_k, dim=1)
            for i, token in enumerate(tokens[:-1]):
                token_id = inputs["input_ids"][0][i].item()
                rank = (predictions.indices[i] == token_id).nonzero()
                rank = rank.item() + 1 if rank.numel() else top_k + 1
                if rank <= 100: green_yellow += 1
                total += 1
        return green_yellow / total if total else 0

    def calculate_perplexity(self, text, model_name="gpt2"):
        tokenizer = GPT2Tokenizer.from_pretrained("gpt2")
        model = GPT2LMHeadModel.from_pretrained("gpt2")
        inputs = tokenizer(text, return_tensors="pt", truncation=True, max_length=1024)
        with torch.no_grad():
            loss = model(**inputs, labels=inputs["input_ids"]).loss
        return torch.exp(loss).item()

    def detect_gpt_free(self, text):
        gltr_score = self.calculate_gltr_score(text)
        roberta = pipeline("text-classification", model="Hello-SimpleAI/chatgpt-detector-roberta")
        roberta_result = roberta(text[:512])[0]
        perplexity = self.calculate_perplexity(text)
        criteria = [
            gltr_score > 0.7,
            roberta_result["label"] == "AI" and roberta_result["score"] > 0.8,
            perplexity < 30
        ]
        return "AI" if sum(criteria) >= 2 else "Human"
    
    def ai_detect(self, text, language="auto", model_name="Hello-SimpleAI/chatgpt-detector-roberta", max_length=512, threshold=0.12):
        try:
            import torch
            import re
            import numpy as np
            from langdetect import detect
            from transformers import pipeline, AutoTokenizer, AutoModelForSequenceClassification, GPT2LMHeadModel
            
            # Deteksi bahasa jika auto
            if language == "auto":
                try:
                    detected_lang = detect(text)
                    print(f"Auto-detected language: {detected_lang}")
                except Exception:
                    detected_lang = "id"  # Default ke Indonesia jika gagal deteksi
                    print("Language detection failed, defaulting to 'id'")
            else:
                detected_lang = language
                print(f"Using specified language: {detected_lang}")
            
            # Membuat cache key berdasarkan text dan bahasa
            import hashlib
            cache_key = hashlib.md5((text[:1000] + str(language)).encode()).hexdigest()
            if hasattr(self, 'ai_detect_cache') and cache_key in self.ai_detect_cache:
                print("Using cached detection result")
                return self.ai_detect_cache[cache_key]
                
            # Khusus untuk Bahasa Indonesia
            if detected_lang == "id" or language == "id":
                print("Processing Indonesian text with specialized approach")
                
                # Buat salinan teks asli
                original_text = text
                
                # Jalankan deteksi pada teks asli (bahasa Indonesia)
                # dan juga terjemahkan ke bahasa Inggris untuk deteksi tambahan
                print("Preparing to translate text for additional analysis")
                
                # 1. Terjemahkan sebagian teks ke bahasa Inggris untuk analisis tambahan
                try:
                    # Terjemahkan hanya sebagian teks (max 2000 karakter) untuk mengurangi overhead
                    sample_text = text[:2000] if len(text) > 2000 else text
                    
                    # Terjemahkan dengan Google Translate
                    try:
                        from googletrans import Translator
                        translator = Translator()
                        translated_text = translator.translate(sample_text, src="id", dest="en").text
                        print("Using Google Translate for analysis support")
                    except Exception as e:
                        print(f"Google Translate failed: {e}, using Mistral Translate")
                        # Fallback ke Mistral translate
                        translated_text = self.mistral_translate(sample_text, src="id", dest="en")
                    
                    print(f"Translation sample (first 100 chars): {translated_text[:100]}")
                except Exception as e:
                    print(f"Translation failed: {e}, proceeding with Indonesian text only")
                    translated_text = None
                    
                # 2. Siapkan model-model deteksi
                try:
                    # Model yang bekerja pada berbagai bahasa
                    roberta_detector = pipeline("text-classification", 
                                              model="Hello-SimpleAI/chatgpt-detector-roberta",
                                              device=0 if torch.cuda.is_available() else -1)
                    
                    # Deteksi khusus bahasa Indonesia jika tersedia
                    try:
                        indo_model = pipeline("text-classification", 
                                             model="indolem/indobert-base-uncased",
                                             device=0 if torch.cuda.is_available() else -1)
                        print("Using IndoBERT model for Indonesian analysis")
                    except Exception as e:
                        print(f"IndoBERT loading failed: {e}")
                        indo_model = None
                        
                except Exception as e:
                    print(f"Model loading failed: {e}")
                    raise e
                
                # Buat chunks dari teks untuk analisis terperinci
                chunks = []
                sentences = re.split(r'(?<=[.!?])\s+', text)
                current_chunk = ""
                
                for sentence in sentences:
                    if len(current_chunk) + len(sentence) < max_length:
                        current_chunk += sentence + " "
                    else:
                        if current_chunk:
                            chunks.append(current_chunk.strip())
                        current_chunk = sentence + " "
                
                if current_chunk:
                    chunks.append(current_chunk.strip())
                
                if not chunks:
                    chunks = [text[:max_length]]
                
                # Hasil analisis per chunk
                chunk_results = []
                
                # Hitung total skor
                total_score1 = 0  # Skor dari roberta detector (untuk teks Indonesia)
                total_score2 = 0  # Skor dari analisis pola bahasa Indonesia
                total_score_translated = 0  # Skor dari teks terjemahan (jika tersedia)
                total_perplexity = 0  # Perplexity
                
                # Buat daftar untuk hasil terjemahan jika tersedia
                translated_chunks = []
                if translated_text:
                    # Kita hanya bisa menganalisis sampel terjemahan
                    translated_chunks = [translated_text]
                
                # Pertama, analisis teks asli bahasa Indonesia
                for i, chunk in enumerate(chunks):
                    print(f"Processing Indonesian chunk {i+1}/{len(chunks)}, length: {len(chunk)}")
                    
                    # 1. Roberta Model (bekerja untuk banyak bahasa)
                    roberta_result = roberta_detector(chunk[:512])[0]
                    score1 = roberta_result["score"] if roberta_result["label"] == "AI" else 1 - roberta_result["score"]
                    
                    # 2. GPT2 Perplexity (kurang optimal untuk bahasa Indonesia tapi masih bisa digunakan)
                    try:
                        perplexity = self.calculate_perplexity(chunk, model_name="gpt2")
                    except Exception as e:
                        print(f"Perplexity calculation failed: {e}")
                        perplexity = 50  # Default value
                    
                    # 3. Tambahan: Analisis pola linguistik khusus bahasa Indonesia
                    # Deteksi pola penggunaan kata penghubung dalam bahasa Indonesia
                    indo_connectors = ["dan", "atau", "tetapi", "namun", "meskipun", "walaupun", "karena", "sebab", 
                                      "oleh", "sehingga", "untuk", "agar", "supaya", "jika", "apabila", "andaikata",
                                      "seandainya", "bahwa", "maka", "dengan", "serta", "lalu", "kemudian", "selanjutnya"]
                    
                    # Deteksi pola repetitif (teks AI cenderung menggunakan struktur kalimat yang sama berulang-ulang)
                    sentences = [s.strip() for s in re.split(r'[.!?]', chunk) if s.strip()]
                    sentence_lengths = [len(s.split()) for s in sentences if s]
                    length_variety = np.std(sentence_lengths) if sentences and len(sentences) > 1 else 0
                    
                    words = chunk.lower().split()
                    connector_count = sum(1 for word in words if word in indo_connectors)
                    
                    # Hitung skor berdasarkan pola linguistik
                    connector_ratio = connector_count / len(words) if words else 0
                    repetition_score = min(1.0, 1 / (length_variety + 1) * 5)  # Makin seragam, makin AI
                    
                    # Gabungkan skor dari pola linguistik
                    score2 = (connector_ratio * 5 + repetition_score) / 2
                    
                    # Gabungkan hasil untuk chunk ini
                    is_ai = bool((score1 * 0.6 + score2 * 0.4) > threshold)  # Konversi ke Python bool standard
                    
                    chunk_result = {
                        "start": text.find(chunk),
                        "end": text.find(chunk) + len(chunk),
                        "score1": float(score1),  # Konversi ke float standard
                        "score2": float(score2),  # Konversi ke float standard
                        "perplexity": float(perplexity),  # Konversi ke float standard
                        "is_ai": is_ai
                    }
                    
                    chunk_results.append(chunk_result)
                    
                    # Update total skor
                    total_score1 += score1
                    total_score2 += score2
                    total_perplexity += perplexity
                
                # Analisis teks terjemahan jika tersedia
                translated_score = 0
                if translated_chunks:
                    print("Analyzing translated text for additional insights")
                    for i, chunk in enumerate(translated_chunks):
                        try:
                            # Gunakan model yang sama untuk deteksi bahasa Inggris
                            trans_result = roberta_detector(chunk[:512])[0]
                            chunk_trans_score = trans_result["score"] if trans_result["label"] == "AI" else 1 - trans_result["score"]
                            translated_score += chunk_trans_score
                            
                            print(f"Translated text AI score: {chunk_trans_score}")
                        except Exception as e:
                            print(f"Error analyzing translated text: {e}")
                    
                    # Rata-rata skor terjemahan
                    if translated_chunks:
                        translated_score /= len(translated_chunks)
                
                # Hitung rata-rata skor dari teks asli
                if chunks:
                    avg_score1 = float(total_score1 / len(chunks))
                    avg_score2 = float(total_score2 / len(chunks))
                    avg_perplexity = float(total_perplexity / len(chunks))
                else:
                    avg_score1 = 0.0
                    avg_score2 = 0.0
                    avg_perplexity = 0.0
                
                # Gabungkan hasil dari teks asli dan terjemahan (jika tersedia)
                # Beri bobot lebih tinggi pada teks asli
                if translated_text:
                    final_score = float(avg_score1 * 0.5 + avg_score2 * 0.3 + translated_score * 0.2)
                    print(f"Combined scores - Original: {avg_score1}, Patterns: {avg_score2}, Translated: {translated_score}")
                else:
                    final_score = float(avg_score1 * 0.7 + avg_score2 * 0.3)
                    print(f"Combined scores - Original: {avg_score1}, Patterns: {avg_score2}")
                
                # Buat objek hasil yang lengkap
                result = {
                    "chunks": chunk_results,
                    "scores": {
                        "model1": float(avg_score1),
                        "model2": float(avg_score2),
                        "perplexity": float(avg_perplexity)
                    },
                    "ai_probability": float(final_score),
                    "is_ai": bool(final_score > threshold),
                    "language": str(detected_lang),
                    "threshold": float(threshold)
                }
                
                # Simpan hasil ke cache untuk penggunaan di masa depan
                if not hasattr(self, 'ai_detect_cache'):
                    self.ai_detect_cache = {}
                self.ai_detect_cache[cache_key] = result
                
                return result
                
            # Untuk bahasa lain (termasuk Inggris)
            else:
                print("Processing non-Indonesian text with standard models")
                # Standard approach for English and other languages
                detector1 = pipeline("text-classification", model=model_name)
                
                # Process the text in chunks if needed
                if len(text) > max_length:
                    chunks = [text[i:i+max_length] for i in range(0, len(text), max_length)]
                else:
                    chunks = [text]
                
                print(f"Processing {len(chunks)} chunks for AI detection")
                
                results = []
                ai_scores = []
                
                for i, chunk in enumerate(chunks):
                    if not chunk.strip():
                        continue
                        
                    # Primary detector
                    try:
                        result1 = detector1(chunk)[0]
                        score1 = result1["score"] if result1["label"] == "AI" else 1 - result1["score"]
                        ai_scores.append(score1)
                        
                        chunk_result = {
                            "text": chunk[:100] + "...", # Tunjukkan awal teks saja untuk mengurangi ukuran respons
                            "ai_score": float(score1),
                            "is_ai": bool(score1 > threshold)
                        }
                        results.append(chunk_result)
                        
                    except Exception as e:
                        print(f"Error in chunk {i}: {e}")
                
                # Calculate overall AI probability
                avg_score = float(sum(ai_scores) / len(ai_scores)) if ai_scores else 0.0
                
                result = {
                    "chunks": results,
                    "ai_probability": float(avg_score),
                    "is_ai": bool(avg_score > threshold),
                    "language": str(detected_lang),
                    "threshold": float(threshold)
                }
                
                # Cache the result
                if not hasattr(self, 'ai_detect_cache'):
                    self.ai_detect_cache = {}
                self.ai_detect_cache[cache_key] = result
                
                return result
                
                # Process English text with standard approach
                # Standard approach for English text using multiple detectors
                detector1 = pipeline("text-classification", model=model_name)
                
                # Try to use a secondary detector if available
                try:
                    detector2 = pipeline("text-classification", model="roberta-base-openai-detector")
                except Exception as e:
                    print(f"Secondary detector not available: {e}")
                    detector2 = None
                
                # Process the text in chunks if needed
                if len(text) > max_length:
                    chunks = [text[i:i+max_length] for i in range(0, len(text), max_length)]
                else:
                    chunks = [text]
                
                print(f"Processing {len(chunks)} chunks for AI detection")
                
                results = []
                ai_scores = []
                
                for i, chunk in enumerate(chunks):
                    if not chunk.strip():
                        continue
                        
                    # Primary detector
                    try:
                        result1 = detector1(chunk)[0]
                        score1 = result1["score"] if result1["label"] == "AI" else 1 - result1["score"]
                        ai_scores.append(score1)
                        
                        chunk_result = {
                            "text": chunk,
                            "ai_score": score1,
                            "is_ai": score1 > threshold
                        }
                        
                        # Add secondary detector result if available
                        if detector2:
                            result2 = detector2(chunk)[0]
                            score2 = result2["score"] if result2["label"] == "AI" else 1 - result2["score"]
                            chunk_result["ai_score2"] = score2
                            chunk_result["is_ai"] = (score1 + score2) / 2 > threshold
                        
                        results.append(chunk_result)
                        
                    except Exception as e:
                        print(f"Error in chunk {i}: {e}")
                
                # Calculate overall AI probability
                avg_score = sum(ai_scores) / len(ai_scores) if ai_scores else 0
                
                return {
                    "chunks": results,
                    "ai_probability": avg_score,
                    "is_ai": avg_score > threshold,
                    "language": detected_lang
                }
                
        except Exception as e:
            print(f"Error in AI detection: {e}")
            import traceback
            traceback.print_exc()
            
            # Fallback to basic detection
            try:
                from transformers import pipeline
                detector = pipeline("text-classification", model="Hello-SimpleAI/chatgpt-detector-roberta")
                
                # Truncate text if needed
                truncated_text = text[:512] if len(text) > 512 else text
                
                # Basic detection
                result = detector(truncated_text)[0]
                score = result["score"] if result["label"] == "AI" else 1 - result["score"]
                
                return {
                    "chunks": [{
                        "text": truncated_text,
                        "ai_score": float(score),
                        "is_ai": bool(score > threshold)
                    }],
                    "ai_probability": float(score),
                    "is_ai": bool(score > threshold),
                    "language": "unknown"
                }
            except Exception as e2:
                print(f"Fallback detection also failed: {e2}")
                return {
                    "chunks": [],
                    "ai_probability": 0.5,  # Inconclusive
                    "is_ai": False,
                    "language": "unknown",
                    "error": str(e),
                    "status": "failed"
                }
                if len(current_chunk) + len(sentence) + 1 < max_length:
                    current_chunk += sentence + " "
                else:
                    chunks.append(current_chunk.strip())
                    current_chunk = sentence + " "
            if current_chunk:
                chunks.append(current_chunk.strip())

            ai_spans = []
            idx = 0
            chunk_results = []
            for chunk in chunks:
                result1 = detector(chunk)[0]
                result2 = gpt_detector(chunk)[0]
                perplexity = self.calculate_perplexity(chunk)
                gltr_score = self.calculate_gltr_score(chunk)
                is_ai = (
                    (result1['label'].lower() == 'ai' and result1['score'] > threshold) or
                    (result2['label'].lower() == 'ai' and result2['score'] > threshold) or
                    (perplexity < 50) or
                    (gltr_score > 0.7)
                )
                chunk_results.append({
                    "start": idx,
                    "end": idx + len(chunk),
                    "label1": result1['label'],
                    "score1": result1['score'],
                    "label2": result2['label'],
                    "score2": result2['score'],
                    "perplexity": perplexity,
                    "gltr_score": gltr_score,
                    "is_ai": is_ai
                })
                if is_ai:
                    ai_spans.append({"start": idx, "end": idx + len(chunk)})
                idx += len(chunk) + 1

            return {
                "ai_spans": ai_spans,
                "chunks": chunk_results
            }
        except Exception as e:
            return {"ai_spans": [], "chunks": [], "error": str(e)}

    def check_similarity(self, text1, text2):
        from sentence_transformers import SentenceTransformer, util
        model = SentenceTransformer('all-MiniLM-L6-v2')
        emb1 = model.encode(text1, convert_to_tensor=True)
        emb2 = model.encode(text2, convert_to_tensor=True)
        similarity = util.pytorch_cos_sim(emb1, emb2)
        return float(similarity[0][0])
    
@app.route("/api/chat", methods=["POST"])
def chat():
    data = request.get_json()
    history = data.get("history", [])

    chatbot = Api()
    reply = chatbot.chat_llama(json.dumps(history))

    return jsonify({"reply": reply})

@app.route('/')
def index():
    return send_from_directory('.', 'DocMaster.html')

@app.route('/simple')
def simple():
    return send_from_directory('.', 'DocMaster_Simple.html')

@app.route('/api/feature/access', methods=['POST'])
def feature_access():
    data = request.get_json()
    email = data.get('email')
    feature = data.get('feature')
    # 1. Cek status premium user di database
    user = get_user_from_db(email)  # Implementasikan fungsi ini
    if not user:
        return jsonify({'allowed': False, 'reason': 'User not found'})
    if user['subscription_type'] == 'premium':
        return jsonify({'allowed': True})
    # 2. Jika non-premium, cek token
    tokens = user.get('tokens', 0)
    if tokens > 0:
        # Kurangi token di database
        update_user_tokens(email, tokens - 1)
        return jsonify({'allowed': True})
    else:
        return jsonify({'allowed': False, 'reason': 'Token habis'})

if __name__ == '__main__':
    api = Api()
    
    # Start Flask in a separate thread
    import threading
    import time
    
    # Find available port
    port = find_available_port()
    
    # Start Flask in thread
    flask_thread = threading.Thread(target=lambda: app.run(host='127.0.0.1', port=port, debug=False))
    flask_thread.daemon = True
    flask_thread.start()
    
    # Wait a moment for Flask to start
    time.sleep(1)
    
    # Start pywebview with main DocMaster application
    window = webview.create_window(
        'DocMaster - Professional Document Processor',
        f'http://127.0.0.1:{port}/',  # Menggunakan route utama, bukan /simple
        width=1200,
        height=800,
        resizable=True,
        js_api=api
    )
    webview.start(debug=True)